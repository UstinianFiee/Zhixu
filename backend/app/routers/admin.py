from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlmodel import Session, select, func
from app.database import get_session
from app.models import WordList, Word, DailyContent, User
from app.schemas import (
    WordListResponse, WordListCreate, WordImportBatch,
    DailyContentResponse, DailyContentCreate, UserResponse,
)
from app.deps import get_admin_user
from app.auth import hash_password
import io

router = APIRouter()


# ── 词库管理 ──
@router.get("/admin/wordlists", response_model=list[WordListResponse])
def admin_get_wordlists(
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    wordlists = db.exec(select(WordList).order_by(WordList.id)).all()
    return [WordListResponse.model_validate(wl) for wl in wordlists]


@router.post("/admin/wordlists", response_model=WordListResponse)
def admin_create_wordlist(
    data: WordListCreate,
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    wl = WordList(name=data.name, description=data.description, is_preset=True)
    db.add(wl)
    db.commit()
    db.refresh(wl)
    return WordListResponse.model_validate(wl)


@router.get("/admin/wordlists/template")
def admin_download_template(admin: User = Depends(get_admin_user)):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "单词模板"
    ws.append(["单词", "释义", "音标(可选)", "例句(可选)"])
    ws.append(["abandon", "放弃", "/əˈbændən/", "Don't abandon hope."])
    ws.append(["ability", "能力", "/əˈbɪləti/", "She has the ability to succeed."])
    ws.append(["abroad", "在国外", "/əˈbrɔːd/", "He studied abroad for two years."])
    ws.column_dimensions["A"].width = 18
    ws.column_dimensions["B"].width = 20
    ws.column_dimensions["C"].width = 18
    ws.column_dimensions["D"].width = 35
    buf = io.BytesIO()
    wb.save(buf)
    buf.seek(0)
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=wordlist_template.xlsx"},
    )


@router.post("/admin/wordlists/{wordlist_id}/words")
def admin_import_words(
    wordlist_id: int,
    data: WordImportBatch,
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    wl = db.get(WordList, wordlist_id)
    if not wl:
        raise HTTPException(status_code=404, detail="词库不存在")

    count = 0
    for item in data.words:
        word = Word(
            wordlist_id=wordlist_id,
            word=item.word,
            translation=item.translation,
            phonetic=item.phonetic,
            example_sentence=item.example_sentence,
        )
        db.add(word)
        count += 1

    wl.word_count += count
    db.add(wl)
    db.commit()
    return {"imported": count, "wordlist_id": wordlist_id}


@router.delete("/admin/wordlists/{wordlist_id}")
def admin_delete_wordlist(
    wordlist_id: int,
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    wl = db.get(WordList, wordlist_id)
    if not wl:
        raise HTTPException(status_code=404, detail="词库不存在")
    # 删除词库下所有单词
    words = db.exec(select(Word).where(Word.wordlist_id == wordlist_id)).all()
    for w in words:
        db.delete(w)
    db.delete(wl)
    db.commit()
    return {"deleted": True}


@router.post("/admin/wordlists/{wordlist_id}/import")
async def admin_import_file(
    wordlist_id: int,
    file: UploadFile = File(...),
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    wl = db.get(WordList, wordlist_id)
    if not wl:
        raise HTTPException(status_code=404, detail="词库不存在")

    filename = file.filename or ""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    content = await file.read()
    words_data = []

    if ext == "txt":
        text = content.decode("utf-8", errors="ignore")
        for line in text.strip().splitlines():
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            parts = line.split(None, 2)
            if len(parts) >= 2:
                words_data.append({
                    "word": parts[0],
                    "translation": parts[1],
                    "phonetic": parts[2] if len(parts) >= 3 else None,
                })

    elif ext in ("xlsx", "xls"):
        from openpyxl import load_workbook
        wb = load_workbook(io.BytesIO(content), read_only=True)
        ws = wb.active
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            if i == 0 and row[0] and str(row[0]).strip().lower() in ("单词", "word"):
                continue  # 跳过表头
            if row[0] and row[1]:
                words_data.append({
                    "word": str(row[0]).strip(),
                    "translation": str(row[1]).strip(),
                    "phonetic": str(row[2]).strip() if len(row) > 2 and row[2] else None,
                    "example_sentence": str(row[3]).strip() if len(row) > 3 and row[3] else None,
                })
        wb.close()

    elif ext == "docx":
        from docx import Document
        doc = Document(io.BytesIO(content))
        for para in doc.paragraphs:
            line = para.text.strip()
            if not line or line.startswith("#"):
                continue
            parts = line.split(None, 2)
            if len(parts) >= 2:
                words_data.append({
                    "word": parts[0],
                    "translation": parts[1],
                    "phonetic": parts[2] if len(parts) >= 3 else None,
                })
        # 也检查表格
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                cells = [c.text.strip() for c in row.cells]
                if i == 0 and cells[0].lower() in ("单词", "word"):
                    continue
                if cells[0] and len(cells) >= 2 and cells[1]:
                    words_data.append({
                        "word": cells[0],
                        "translation": cells[1],
                        "phonetic": cells[2] if len(cells) >= 3 and cells[2] else None,
                        "example_sentence": cells[3] if len(cells) >= 4 and cells[3] else None,
                    })
    else:
        raise HTTPException(status_code=400, detail="不支持的文件格式，请使用 txt / xlsx / docx")

    if not words_data:
        raise HTTPException(status_code=400, detail="文件中未找到有效单词数据")

    count = 0
    for item in words_data:
        word = Word(
            wordlist_id=wordlist_id,
            word=item["word"],
            translation=item["translation"],
            phonetic=item.get("phonetic"),
            example_sentence=item.get("example_sentence"),
        )
        db.add(word)
        count += 1

    wl.word_count += count
    db.add(wl)
    db.commit()
    return {"imported": count, "wordlist_id": wordlist_id}


# ── 文艺内容管理 ──
@router.get("/admin/content", response_model=list[DailyContentResponse])
def admin_get_content(
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    contents = db.exec(
        select(DailyContent).order_by(DailyContent.date.desc())
    ).all()
    return [DailyContentResponse.model_validate(c) for c in contents]


@router.post("/admin/content", response_model=DailyContentResponse)
def admin_create_content(
    data: DailyContentCreate,
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    content = DailyContent(**data.model_dump())
    db.add(content)
    db.commit()
    db.refresh(content)
    return DailyContentResponse.model_validate(content)


@router.put("/admin/content/{content_id}", response_model=DailyContentResponse)
def admin_update_content(
    content_id: int,
    data: DailyContentCreate,
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    content = db.get(DailyContent, content_id)
    if not content:
        raise HTTPException(status_code=404, detail="内容不存在")
    for key, value in data.model_dump().items():
        setattr(content, key, value)
    db.add(content)
    db.commit()
    db.refresh(content)
    return DailyContentResponse.model_validate(content)


@router.delete("/admin/content/{content_id}")
def admin_delete_content(
    content_id: int,
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    content = db.get(DailyContent, content_id)
    if not content:
        raise HTTPException(status_code=404, detail="内容不存在")
    db.delete(content)
    db.commit()
    return {"deleted": True}


# ── 用户管理 ──
@router.get("/admin/users", response_model=list[UserResponse])
def admin_get_users(
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    users = db.exec(select(User).order_by(User.id)).all()
    return [UserResponse.model_validate(u) for u in users]


from pydantic import BaseModel as PydanticBaseModel

class RoleUpdate(PydanticBaseModel):
    is_admin: bool

class PasswordReset(PydanticBaseModel):
    password: str


@router.put("/admin/users/{user_id}/role")
def admin_update_user_role(
    user_id: int,
    data: RoleUpdate,
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    user = db.get(User, user_id)
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    user.is_admin = data.is_admin
    db.add(user)
    db.commit()
    return {"updated": True, "is_admin": user.is_admin}


@router.put("/admin/users/{user_id}/password")
def admin_reset_password(
    user_id: int,
    data: PasswordReset,
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    user = db.get(User, user_id)
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    if len(data.password) < 6:
        raise HTTPException(status_code=400, detail="密码至少6个字符")
    user.password_hash = hash_password(data.password)
    db.add(user)
    db.commit()
    return {"updated": True}


@router.delete("/admin/users/{user_id}")
def admin_delete_user(
    user_id: int,
    admin: User = Depends(get_admin_user),
    db: Session = Depends(get_session),
):
    user = db.get(User, user_id)
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    if user.id == admin.id:
        raise HTTPException(status_code=400, detail="不能删除自己的账号")
    db.delete(user)
    db.commit()
    return {"deleted": True}
